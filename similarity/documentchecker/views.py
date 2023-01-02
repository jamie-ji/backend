from django.shortcuts import render
import urllib.parse
import os
import re
from dateutil import parser
from django.core.files.storage import default_storage
from rest_framework import generics, permissions, status
from rest_framework.response import Response
from docx import Document
from .models import File, Threshold,SimilarityCheck
from .tasks import similaritycheck
from django.db.models import Sum

# Create your views here.

def build_data(queryset):
    
    threshold_obj = Threshold.objects.filter(active=True).first()
    
    data_dict = {}
    
    years = [ year[0].year for year in queryset.values_list('year')]
    distinct_years = list(set(years))
    print('dist year<<<<<<<<<<<',distinct_years)
    
    total_words = queryset.aggregate(total_words=Sum('word_count'))
    
    if len(distinct_years) != int(threshold_obj.distinct_year):
        data_dict['years_status'] = 'More years needed'
    else:
        data_dict['years_status'] = 'pass'
    
    data_dict['total_words'] = total_words['total_words']
    data_dict['total_files'] = queryset.values_list('file').count()
    
    for year in distinct_years:
        
        data_dict[year] = {}
        
        year_querset=queryset.filter(created_at__year=year)
        file_per_year=year_querset.values_list('file').count()
        data_dict[year]['file_count'] = file_per_year
        if file_per_year < int(threshold_obj.file_per_year):
            data_dict[year]['files_status'] = 'More file needed'
        else:
            data_dict[year]['files_status'] = 'pass'
            
        word_count_per_year = year_querset.aggregate(word_count=Sum('word_count'))
        data_dict[year]['word_count'] = word_count_per_year['word_count']
        if word_count_per_year['word_count'] < int(threshold_obj.words_per_year):
            data_dict[year]['words_status'] = 'More words needed'
        else:
            data_dict[year]['words_status'] = 'pass'
    
    return data_dict


def words_count(file):
    
    doc=Document(file)
    prop = doc.core_properties
    # print(dir(prop))
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    text='\n'.join(fullText)
    rj1=re.sub(r'^$\n', '',text, flags=re.MULTILINE)
    return(len(re.findall(r'\w+',rj1)))


class UploadFile(generics.GenericAPIView):
    authentication_classes = []
    permission_classes = []

    def post(self, request):
        
        file = request.data.get("file", None)
        print("file>>>>>>>>>>>>>>>>>",file.name)
        if file is None or file=='':
            return Response({"file": "file is required"}, status=status.HTTP_400_BAD_REQUEST)
        
        extension = os.path.splitext(str(file))[-1].lower()
        print("extention>>>>>>>>>>>>>>>>>>",extension)
        if extension == ".docx":
            try:
                doc = Document(file)
            except Exception as e:
                return Response({"file":"File is Corrupted or ".format(e)},status=status.HTTP_400_BAD_REQUEST)
            prop = doc.core_properties
            
            author=prop.author
            year=prop.created
            word_count=words_count(file)
            print("year>>>>>>>>>>>",year)
            print("year>>>>>>>>>>>",type(year))
            datetime = parser.parse(str(year))
            if year is None or year=='':
                return Response({"Year": "File Without Year Not Allowed {}".format(file)}, status=status.HTTP_400_BAD_REQUEST)
            if author is None or author=='':
                return Response({"Author": "{} does not have author ".format(file)}, status=status.HTTP_400_BAD_REQUEST)
            full_path = "documents/" + urllib.parse.quote(file.name.replace(" ", "-"), safe="")
            
            file_obj= File()
            file = default_storage.save(full_path,file)
            file_obj.author=author
            file_obj.created_at=datetime
            file_obj.file=file
            file_obj.word_count=word_count
            file_obj.save()
            
            url = default_storage.url(file)
            
            data_dict={"full_url": request.build_absolute_uri(url), 
                       "path": file_obj.file.name,
                       'author':file_obj.author,
                       'id': file_obj.id,
                       "created_at":file_obj.created_at,
                       'word_count':file_obj.word_count
                       }
            
            return Response(data_dict, status=status.HTTP_200_OK)
                
        # elif ext==".doc":
        #     print("hererere")
        #     docfile.delay(file=str(file))
        #     return Response("File Converted")
        else:
            return Response("File Not Valid", status=status.HTTP_400_BAD_REQUEST)
        
class DocumentCheck(generics.GenericAPIView):
    authentication_classes = []
    permission_classes = []
    def post(self, request):
        file = request.data.get("file_id")
        if file is None or file==[]:
            return Response("Please Select the File",status=status.HTTP_400_BAD_REQUEST)
        author = request.data.get("author")
        if author is None or author=="":
            return Response("Please Select the Author",status=status.HTTP_400_BAD_REQUEST)
        print('id',author)
        # task_obj = Task()
        # task_obj.status = 'Files Submitted'
        # task_obj.save()
        # print("obj>>>>>>>",task_obj)
        
        similaritycheck.delay(file=file,author=author)
        # similaritycheck.delay(file=file)
        
        
        return Response("Files submitted", status=status.HTTP_200_OK)

    def get(self, request, *args, **kwargs):
        author =self.request.query_params.get("author", None)
        print(author)
        if author:
            queryset = File.objects.filter(author=author) 
            data_dict = build_data(queryset)
            if data_dict:
                return Response(data_dict, status=status.HTTP_200_OK)
            else:
                return Response({'data': 'data not found'}, status.HTTP_400_BAD_REQUEST)
        else:
            return Response({'author': 'author not provided'}, author)



